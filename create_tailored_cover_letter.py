"""
Tailored Cover Letter Generator
Generates deeply customized cover letters for cybersecurity program management positions.

Input Options:
1. Job URL (LinkedIn, company career page)
2. Job Document Upload (PDF, Word .docx, or .txt file)

Output:
- Professionally formatted Word document (.docx)
- Company-specific research and customization
- Saved to TAILORED_COVER_LETTERS directory
"""

import os
import re
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Try importing PDF libraries
try:
    import pdfplumber
    PDF_LIBRARY = 'pdfplumber'
except ImportError:
    try:
        import PyPDF2
        PDF_LIBRARY = 'PyPDF2'
    except ImportError:
        PDF_LIBRARY = None
        print("WARNING: No PDF library installed. PDF uploads will not work.")
        print("Install with: pip install pdfplumber")

# Try importing Playwright for URL scraping
try:
    from playwright.sync_api import sync_playwright
    PLAYWRIGHT_AVAILABLE = True
except ImportError:
    PLAYWRIGHT_AVAILABLE = False
    print("WARNING: Playwright not installed. URL extraction will not work.")
    print("Install with: pip install playwright && python -m playwright install chromium")


# =============================================================================
# CONFIGURATION
# =============================================================================

# User selects input method
INPUT_METHOD = "url"  # Options: "url" or "document"

# Job URL (if INPUT_METHOD = "url")
JOB_URL = "https://jpmc.fa.oraclecloud.com/hcmUI/CandidateExperience/en/sites/CX_1001/job/210559736"

# Job Document Path (if INPUT_METHOD = "document")
JOB_DOCUMENT_PATH = "C:/Users/derri/sample_job_description.txt"

# Customization Options
TONE = "Professional"  # Options: Professional, Enthusiastic, Strategic, Technical
LENGTH = "Standard"    # Options: Concise (3 paragraphs), Standard (4), Detailed (5)
FOCUS = "Program Management"  # Options: Leadership, Technical, Program Management, Cross-functional

# Output Directory
OUTPUT_DIR = "C:/Users/derri/TAILORED_COVER_LETTERS"

# User Contact Information
USER_INFO = {
    "name": "Justin Washington",
    "email": "justinwashington@gmail.com",
    "phone": "(713) 927-5607",
    "linkedin": "linkedin.com/in/justintwashington",
    "location": "Houston, TX"
}


# =============================================================================
# FILE EXTRACTION FUNCTIONS
# =============================================================================

def extract_pdf(file_path):
    """Extract text from PDF file using available library"""
    if PDF_LIBRARY == 'pdfplumber':
        return extract_pdf_pdfplumber(file_path)
    elif PDF_LIBRARY == 'PyPDF2':
        return extract_pdf_pypdf2(file_path)
    else:
        raise ImportError("No PDF library installed. Install pdfplumber: pip install pdfplumber")


def extract_pdf_pdfplumber(file_path):
    """Extract text from PDF using pdfplumber (more reliable)"""
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text


def extract_pdf_pypdf2(file_path):
    """Extract text from PDF using PyPDF2 (fallback)"""
    text = ""
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    return text


def extract_docx(file_path):
    """Extract text from Word document"""
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


def extract_txt(file_path):
    """Extract text from plain text file"""
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()


def detect_file_type(file_path):
    """Detect file type and extract text accordingly"""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        print(f"[PDF] Extracting: {os.path.basename(file_path)}")
        return extract_pdf(file_path)
    elif ext == '.docx':
        print(f"[DOCX] Extracting Word document: {os.path.basename(file_path)}")
        return extract_docx(file_path)
    elif ext == '.txt':
        print(f"[TXT] Reading text file: {os.path.basename(file_path)}")
        return extract_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx, .txt")


# =============================================================================
# URL EXTRACTION FUNCTIONS
# =============================================================================

def extract_job_from_url(url):
    """Extract job details from URL using Playwright"""
    if not PLAYWRIGHT_AVAILABLE:
        raise ImportError("Playwright not installed. Install: pip install playwright")

    print(f"[WEB] Fetching job from URL: {url}")

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()

        try:
            page.goto(url, timeout=30000)
            page.wait_for_timeout(5000)  # Wait longer for JavaScript to load

            # Try to wait for job description content
            try:
                page.wait_for_selector('article, .job-description, .description', timeout=5000)
            except:
                pass  # Content might be structured differently

            # Extract full page text
            text = page.inner_text('body')

            browser.close()
            return text

        except Exception as e:
            browser.close()
            raise Exception(f"Failed to extract job from URL: {str(e)}")


# =============================================================================
# JOB PARSING FUNCTIONS
# =============================================================================

def detect_company_from_url(url):
    """Detect company name from URL domain"""
    url_lower = url.lower()

    if 'jpmc' in url_lower or 'jpmorganchase' in url_lower:
        return "JPMorgan Chase"
    elif 'oracle' in url_lower:
        return "Oracle"
    elif 'microsoft' in url_lower:
        return "Microsoft"
    elif 'google' in url_lower:
        return "Google"
    elif 'amazon' in url_lower:
        return "Amazon"
    elif 'apple' in url_lower:
        return "Apple"
    elif 'meta' in url_lower or 'facebook' in url_lower:
        return "Meta"
    elif 'linkedin' in url_lower:
        # LinkedIn job board - company might be in path
        return None
    else:
        return None


def parse_job_description(text, source_url=None):
    """Parse job description text to extract key details"""

    job_data = {
        "company": "",
        "title": "",
        "location": "",
        "salary": "",
        "description": text,
        "requirements": [],
        "responsibilities": [],
        "source_url": source_url
    }

    # Split into lines for better parsing
    lines = text.split('\n')
    first_lines = [line.strip() for line in lines[:10] if line.strip()]

    # Extract company name (heuristic: often in first 200 chars or after "Company:" or "@")
    company_patterns = [
        r"(?:Company|Organization):\s*([A-Z][A-Za-z\s&.]+)",
        r"^([A-Z][A-Za-z\s&.]+Corporation)$",  # Company with Corporation
        r"^([A-Z][A-Za-z\s&.]+Inc\.)$",  # Company with Inc.
        r"^([A-Z][A-Za-z\s&.]+LLC)$",  # Company with LLC
        r"^(Microsoft|Google|Amazon|Apple|Oracle|IBM|Meta|Tesla|JPMorgan|Goldman Sachs)$",  # Known companies
        r"([A-Z][A-Za-z\s&.]+)\s*(?:is seeking|is hiring)",
        r"Join\s+([A-Z][A-Za-z\s&.]+)",
    ]

    for line in first_lines:
        for pattern in company_patterns:
            match = re.search(pattern, line)
            if match:
                job_data["company"] = match.group(1).strip()
                break
        if job_data["company"]:
            break

    # Extract job title (usually first line)
    title_patterns = [
        r"(?:Position|Title|Role):\s*([A-Z][A-Za-z\s\-/&]+)",
        r"^([A-Z][A-Za-z\s\-/&]{15,80})$",  # Standalone line with title-like text (longer = likely a title)
    ]

    for line in first_lines:
        # Skip if this is the company line
        if line == job_data.get("company"):
            continue
        for pattern in title_patterns:
            match = re.search(pattern, line)
            if match:
                job_data["title"] = match.group(1).strip()
                break
        if job_data["title"]:
            break

    # If title still not found, assume first non-empty line is the title
    if not job_data["title"] and first_lines:
        potential_title = first_lines[0]
        if potential_title != job_data.get("company") and len(potential_title) > 10:
            job_data["title"] = potential_title

    # Extract location
    location_patterns = [
        r"(?:Location|Where):\s*([A-Za-z\s,]+(?:TX|NY|CA|Remote|Hybrid))",
        r"((?:[A-Z][a-z]+,\s*[A-Z]{2}|Remote|Hybrid))",
    ]
    for pattern in location_patterns:
        match = re.search(pattern, text[:1000])
        if match:
            job_data["location"] = match.group(1).strip()
            break

    # Extract salary (if present)
    salary_patterns = [
        r"\$[\d,]+\s*-\s*\$[\d,]+",
        r"(?:Salary|Compensation):\s*\$[\d,]+",
    ]
    for pattern in salary_patterns:
        match = re.search(pattern, text)
        if match:
            job_data["salary"] = match.group(0).strip()
            break

    # Extract requirements (look for bullet points or "Required:" section)
    requirements_section = re.search(
        r"(?:Required|Requirements|Qualifications|Must Have):(.+?)(?:Preferred|Responsibilities|About|$)",
        text,
        re.DOTALL | re.IGNORECASE
    )
    if requirements_section:
        req_text = requirements_section.group(1)
        # Find bullet points or numbered lists
        bullets = re.findall(r"(?:^|\n)\s*(?:[-•*]|\d+\.)\s*(.+)", req_text)
        job_data["requirements"] = [b.strip() for b in bullets if len(b.strip()) > 10]

    # Extract responsibilities
    resp_section = re.search(
        r"(?:Responsibilities|You Will|Your Role|What You'll Do):(.+?)(?:Required|Qualifications|About|$)",
        text,
        re.DOTALL | re.IGNORECASE
    )
    if resp_section:
        resp_text = resp_section.group(1)
        bullets = re.findall(r"(?:^|\n)\s*(?:[-•*]|\d+\.)\s*(.+)", resp_text)
        job_data["responsibilities"] = [b.strip() for b in bullets if len(b.strip()) > 10]

    return job_data


# =============================================================================
# COMPANY RESEARCH (Simplified - in production, use web search APIs)
# =============================================================================

def research_company(company_name):
    """
    Research company using 5-step framework from claude.md
    NOTE: This is a simplified version. In production, implement web scraping/API calls
    """

    # Placeholder company research data
    # In production, implement web scraping to gather:
    # 1. Mission & Values
    # 2. Industry Initiatives
    # 3. Security Team Culture
    # 4. Compliance Environment

    company_db = {
        "JPMorgan": {
            "formal_name": "JPMorgan Chase & Co.",
            "mission": "enable business by keeping the firm safe, stable, and resilient",
            "values": ["integrity", "excellence", "innovation"],
            "initiatives": ["$1.5 trillion Security & Resiliency Initiative", "AI security", "quantum computing"],
            "color": "#003366",  # Navy blue
            "industry": "Financial Services"
        },
        "Oracle": {
            "formal_name": "Oracle Corporation",
            "mission": "help people see data in new ways, discover insights, and unlock endless possibilities",
            "values": ["integrity", "innovation", "customer satisfaction"],
            "initiatives": ["Oracle Health Federal", "Cloud security", "Healthcare IT modernization"],
            "color": "#FF0000",  # Oracle red
            "industry": "Technology / Healthcare"
        },
        "Microsoft": {
            "formal_name": "Microsoft Corporation",
            "mission": "empower every person and organization on the planet to achieve more",
            "values": ["growth mindset", "customer obsession", "diversity & inclusion"],
            "initiatives": ["Azure security", "AI integration", "Zero Trust architecture"],
            "color": "#0078D4",  # Microsoft blue
            "industry": "Technology"
        }
    }

    # Match company name (case-insensitive, partial match)
    for key, data in company_db.items():
        if key.lower() in company_name.lower():
            return data

    # Default company research if not in database
    return {
        "formal_name": company_name,
        "mission": "deliver excellence and innovation",
        "values": ["integrity", "innovation", "customer focus"],
        "initiatives": ["cybersecurity modernization", "cloud transformation"],
        "color": "#003366",  # Default navy blue
        "industry": "Technology"
    }


# =============================================================================
# COVER LETTER GENERATION
# =============================================================================

def generate_opening_paragraph(job_data, company_research, tone):
    """Generate compelling opening paragraph"""

    company = company_research["formal_name"]
    title = job_data["title"]
    initiative = company_research["initiatives"][0] if company_research["initiatives"] else "cybersecurity initiatives"

    if tone == "Enthusiastic":
        opening = (
            f"I am excited to apply for the {title} position at {company}. "
            f"With over 10 years of experience driving cybersecurity program management in enterprise and federal environments, "
            f"I am particularly drawn to {company}'s commitment to {initiative}. "
            f"My track record of building risk governance frameworks, managing cross-functional stakeholder relationships, "
            f"and delivering measurable security outcomes aligns directly with this opportunity to contribute to your team."
        )
    elif tone == "Strategic":
        opening = (
            f"As a strategic cybersecurity program leader with 10+ years of experience, I am writing to express my strong interest "
            f"in the {title} role at {company}. Your organization's focus on {initiative} represents exactly the kind of "
            f"forward-thinking security approach where my expertise in risk frameworks, compliance program management, "
            f"and stakeholder engagement can drive meaningful impact."
        )
    elif tone == "Technical":
        opening = (
            f"I am applying for the {title} position at {company} with a strong background in cybersecurity program management, "
            f"vulnerability remediation, and compliance frameworks including NIST, ISO 27001, and CVSS scoring methodologies. "
            f"My technical expertise combined with proven program leadership aligns well with {company}'s {initiative}, "
            f"and I am confident I can contribute to your security objectives from day one."
        )
    else:  # Professional (default)
        opening = (
            f"I am writing to apply for the {title} position at {company}. "
            f"With over 10 years of experience leading cybersecurity programs in complex enterprise environments, "
            f"I have developed deep expertise in risk management, vulnerability remediation, and compliance frameworks. "
            f"I am particularly impressed by {company}'s {initiative} and believe my background in building security governance structures "
            f"and managing cross-functional teams positions me well to contribute to your mission."
        )

    return opening


def generate_body_paragraphs(job_data, company_research, focus):
    """Generate 2-3 body paragraphs demonstrating qualifications"""

    paragraphs = []

    # Paragraph 1: Relevant Experience Mapping
    if focus == "Leadership":
        para1 = (
            "Throughout my career, I have led cybersecurity programs for diverse organizations, managing teams, "
            "coordinating with executive stakeholders, and driving security strategy aligned with business objectives. "
            "As Senior Manager of Cybersecurity at Heirclark, I built and operationalized a risk governance framework "
            "that reduced security escalations by 17% through proactive threat monitoring, vulnerability management workflows, "
            "and executive cyber-risk reporting. I have extensive experience translating complex technical requirements "
            "into actionable plans that both technical teams and business leaders can understand and support."
        )
    elif focus == "Technical":
        para1 = (
            "My technical expertise spans vulnerability management, NIST CVSS scoring, compliance frameworks (HIPAA, FedRAMP, ISO 27001), "
            "and security tooling across cloud and on-premises environments. At Heirclark, I managed vulnerability assessment programs "
            "for 40+ enterprise implementations, coordinating scan scheduling, prioritizing findings based on CVSS scoring, "
            "and ensuring timely mitigation of critical and high-severity vulnerabilities with 95%+ SLA compliance. "
            "I have hands-on experience with security tools including Splunk, Tenable, Qualys, ServiceNow, and cloud security platforms."
        )
    elif focus == "Cross-functional":
        para1 = (
            "One of my core strengths is bridging technical security teams with business stakeholders to drive collaborative outcomes. "
            "In my role at Heirclark, I managed relationships across IT, compliance, legal, and business units to deliver unified "
            "security programs that met both technical requirements and business objectives. I facilitated regular risk committee meetings, "
            "translated security findings into business-relevant reports, and built consensus across diverse stakeholder groups "
            "to implement security controls without impeding operational efficiency."
        )
    else:  # Program Management (default)
        para1 = (
            "As a cybersecurity program manager, I excel at breaking down complex security initiatives into executable roadmaps, "
            "coordinating cross-functional teams, and delivering measurable results. At Heirclark, I managed a cybersecurity project "
            "portfolio overseeing vulnerability assessment, remediation tracking, and compliance delivery for 40+ enterprise implementations. "
            "I built PMO structures, established work assignment workflows, and drove timely vulnerability remediation that achieved "
            "95% SLA compliance for critical and high-severity findings. My approach combines agile methodology, risk-based prioritization, "
            "and stakeholder communication to ensure security programs deliver both protection and business value."
        )

    paragraphs.append(para1)

    # Paragraph 2: Measurable Achievements
    para2 = (
        "My contributions have consistently delivered measurable security improvements and business value. "
        "In addition to reducing security escalations by 17%, I have led programs that achieved 100% compliance during regulatory audits, "
        "reduced mean time to remediation by 23%, and built threat intelligence capabilities that identified and mitigated "
        "emerging threats before they impacted production systems. At Deloitte, I delivered cybersecurity advisory services for "
        "100+ healthcare clients, achieving a 93% Net Promoter Score and generating $25M in managed services revenue. "
        "These results demonstrate my ability to deliver both security outcomes and organizational value."
    )

    paragraphs.append(para2)

    return paragraphs


def generate_closing_paragraph(job_data, company_research):
    """Generate mission-aligned closing paragraph"""

    company = company_research["formal_name"]
    mission = company_research["mission"]

    closing = (
        f"I am committed to {company}'s mission to {mission}, and I am confident that my background in cybersecurity program management, "
        f"risk frameworks, and cross-functional leadership will enable me to contribute meaningfully to your team. "
        f"I would welcome the opportunity to discuss how my experience aligns with your needs and how I can help advance your security objectives. "
        f"Thank you for considering my application. I look forward to the possibility of contributing to {company}'s continued success."
    )

    return closing


# =============================================================================
# WORD DOCUMENT CREATION
# =============================================================================

def add_hyperlink(paragraph, url, text, color='0563C1'):
    """Add clickable hyperlink to Word document"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)

    # Set underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def create_cover_letter_document(job_data, company_research, content_paragraphs, output_path):
    """Create formatted Word document"""

    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Header: User contact information
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(USER_INFO["name"])
    run.font.size = Pt(14)
    run.font.bold = True

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = f'{USER_INFO["email"]} | {USER_INFO["phone"]} | {USER_INFO["linkedin"]} | {USER_INFO["location"]}'
    run = contact.add_run(contact_text)
    run.font.size = Pt(10)

    doc.add_paragraph()  # Spacing

    # Date
    date_para = doc.add_paragraph()
    date_para.add_run(datetime.now().strftime("%B %d, %Y"))

    doc.add_paragraph()  # Spacing

    # Company address (if known)
    company_para = doc.add_paragraph()
    company_para.add_run(f"Hiring Manager\n{company_research['formal_name']}\n")
    if job_data.get("location"):
        company_para.add_run(f"{job_data['location']}\n")

    doc.add_paragraph()  # Spacing

    # Salutation
    salutation = doc.add_paragraph()
    salutation.add_run("Dear Hiring Manager,")

    doc.add_paragraph()  # Spacing

    # Add content paragraphs (opening, body, closing)
    for para_text in content_paragraphs:
        para = doc.add_paragraph()
        para.add_run(para_text)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()  # Spacing between paragraphs

    # Closing
    closing = doc.add_paragraph()
    closing.add_run("Sincerely,")

    doc.add_paragraph()  # Spacing

    signature = doc.add_paragraph()
    run = signature.add_run(USER_INFO["name"])
    run.font.bold = True

    # Save document
    doc.save(output_path)
    print(f"[SUCCESS] Cover letter saved: {output_path}")


# =============================================================================
# MAIN EXECUTION
# =============================================================================

def main():
    """Main execution function"""

    print("\n" + "="*70)
    print("  TAILORED COVER LETTER GENERATOR")
    print("="*70 + "\n")

    # Step 1: Get job description text
    if INPUT_METHOD == "url":
        print("[INPUT] Method: Job URL")
        job_text = extract_job_from_url(JOB_URL)
    elif INPUT_METHOD == "document":
        print("[INPUT] Method: Uploaded Document")
        job_text = detect_file_type(JOB_DOCUMENT_PATH)
    else:
        raise ValueError(f"Invalid INPUT_METHOD: {INPUT_METHOD}. Must be 'url' or 'document'")

    print(f"[SUCCESS] Job description extracted ({len(job_text)} characters)\n")

    # Step 2: Parse job description
    print("[PARSE] Parsing job details...")
    source_url = JOB_URL if INPUT_METHOD == "url" else None
    job_data = parse_job_description(job_text, source_url)

    print(f"   Company: {job_data['company'] or 'Not detected'}")
    print(f"   Title: {job_data['title'] or 'Not detected'}")
    print(f"   Location: {job_data['location'] or 'Not detected'}")
    print(f"   Requirements: {len(job_data['requirements'])} found")
    print(f"   Responsibilities: {len(job_data['responsibilities'])} found\n")

    # If company/title not auto-detected, try URL detection or use defaults
    if not job_data["company"]:
        print("[WARNING] Company name not auto-detected.")
        if source_url:
            detected_company = detect_company_from_url(source_url)
            if detected_company:
                job_data["company"] = detected_company
                print(f"[INFO] Detected company from URL: {detected_company}")
            else:
                job_data["company"] = "Target Company"
                print("[INFO] Using default: 'Target Company'")
        else:
            job_data["company"] = "Target Company"
            print("[INFO] Using default: 'Target Company'")

    if not job_data["title"]:
        print("[WARNING] Job title not auto-detected.")
        job_data["title"] = "Cybersecurity Program Manager"
        print("[INFO] Using default: 'Cybersecurity Program Manager'")

    # Step 3: Research company
    print(f"[RESEARCH] Researching {job_data['company']}...")
    company_research = research_company(job_data["company"])
    print(f"   Mission: {company_research['mission']}")
    print(f"   Industry: {company_research['industry']}")
    print(f"   Initiatives: {', '.join(company_research['initiatives'][:2])}\n")

    # Step 4: Generate cover letter content
    print(f"[GENERATE] Creating cover letter...")
    print(f"   Tone: {TONE}")
    print(f"   Length: {LENGTH}")
    print(f"   Focus: {FOCUS}\n")

    opening = generate_opening_paragraph(job_data, company_research, TONE)
    body_paras = generate_body_paragraphs(job_data, company_research, FOCUS)
    closing = generate_closing_paragraph(job_data, company_research)

    # Combine all paragraphs
    if LENGTH == "Concise":
        content_paragraphs = [opening, body_paras[0], closing]
    elif LENGTH == "Detailed":
        content_paragraphs = [opening] + body_paras + [closing]
    else:  # Standard
        content_paragraphs = [opening] + body_paras + [closing]

    # Step 5: Create Word document
    print("[DOCX] Creating Word document...")

    # Create output directory if it doesn't exist
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Generate filename (sanitize by removing invalid characters)
    company_clean = re.sub(r'[^\w\s-]', '', job_data["company"]).strip()
    company_clean = re.sub(r'\s+', '', company_clean)[:20]  # Remove all whitespace

    title_clean = re.sub(r'[^\w\s-]', '', job_data["title"]).strip()
    title_clean = re.sub(r'\s+', '_', title_clean)[:30]  # Replace whitespace with underscore

    filename = f"CoverLetter_{company_clean}_{title_clean}.docx"
    output_path = os.path.join(OUTPUT_DIR, filename)

    create_cover_letter_document(job_data, company_research, content_paragraphs, output_path)

    print("\n" + "="*70)
    print("[SUCCESS] COVER LETTER GENERATION COMPLETE!")
    print("="*70)
    print(f"\n[OUTPUT] Location: {output_path}\n")

    # Display preview
    print("[PREVIEW]\n")
    print("-" * 70)
    for i, para in enumerate(content_paragraphs, 1):
        print(f"\nParagraph {i}:")
        print(para[:200] + "..." if len(para) > 200 else para)
    print("-" * 70 + "\n")


if __name__ == "__main__":
    main()
